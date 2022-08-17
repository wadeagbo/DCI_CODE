import sys, os, time, datetime
import docx
from colorama             import *
from art                  import *
from docx.enum.text       import WD_ALIGN_PARAGRAPH
from docx.shared          import RGBColor, Pt

# initialize color effects of colorama module
init()

# default blank spaces
clear      = "\033[2J\033[1;1f"
space      = "\n" * 20

# time variables
today      = datetime.date.today().strftime("%d-%m-%Y")
now        = str(datetime.datetime.today().strftime("%H:%M:%S %d-%m-%Y"))

here = os.getcwd()
# audio elements
try:
       import vlc
       # audio elements - effect
       info    = vlc.MediaPlayer("Components\\Audio\\Effects\\info.mp3")
       info2   = vlc.MediaPlayer("Components\\Audio\\Effects\\info2.mp3")
       good    = vlc.MediaPlayer("Components\\Audio\\Effects\\good.mp3")
       good2   = vlc.MediaPlayer("Components\\Audio\\Effects\\good2.mp3")
       good3   = vlc.MediaPlayer("Components\\Audio\\Effects\\good3.mp3")
       bad     = vlc.MediaPlayer("Components\\Audio\\Effects\\bad.mp3")
       alarm   = vlc.MediaPlayer("Components\\Audio\\Effects\\alarm.mp3")
       loading = vlc.MediaPlayer("Components\\Audio\\Effects\\loading.mp3")
       effect1 = vlc.MediaPlayer("Components\\Audio\\Effects\\effect1.mp3")
       effect2 = vlc.MediaPlayer("Components\\Audio\\Effects\\effect2.mp3")
       # audio elements - voice
       clearreq     = vlc.MediaPlayer("Components\\Audio\\Voice\\clearreq.mp3")
       clearfail    = vlc.MediaPlayer("Components\\Audio\\Voice\\clearfail.mp3")
       clearconf    = vlc.MediaPlayer("Components\\Audio\\Voice\\clearconf.mp3")
       connected    = vlc.MediaPlayer("Components\\Audio\\Voice\\connected.mp3")
       noconnection = vlc.MediaPlayer("Components\\Audio\\Voice\\noconnection.mp3")
       readytoadd   = vlc.MediaPlayer("Components\\Audio\\Voice\\readytoadd.mp3")
       # audio elements - greeting
       righnow    = datetime.datetime.now()
       if 5 <= righnow.hour < 12:
              greet    = "morning"
              greeting = vlc.MediaPlayer("Components\\Audio\\Voice\\morning.mp3")               
       elif 12 <= righnow.hour < 18:                                                   
              greet    = "afternoon"                                                 
              greeting = vlc.MediaPlayer("Components\\Audio\\Voice\\afternoon.mp3")     
       else:                                                                      
              greet    = "evening"                                                    
              greeting = vlc.MediaPlayer("Components\\Audio\\Voice\\evening.mp3")
except Exception as error:
       print("VLC IMPORT ERROR ("+str(error)+")")
       class NOvlc():
              def play(self):
                     time.sleep(0.001)
              def stop(self):
                     time.sleep(0.001)
       # audio elements - effect
       info    = NOvlc()
       info2   = NOvlc()
       good    = NOvlc()
       good2   = NOvlc()
       good3   = NOvlc()
       bad     = NOvlc()
       alarm   = NOvlc()
       loading = NOvlc()
       effect1 = NOvlc()
       effect2 = NOvlc()
       # audio elements - voice
       clearreq     = NOvlc()
       clearfail    = NOvlc()
       clearconf    = NOvlc()
       connected    = NOvlc()
       noconnection = NOvlc()
       readytoadd   = NOvlc()
       # audio elements - greeting
       righnow    = datetime.datetime.now()
       if 5 <= righnow.hour < 12:
              greet    = "morning"
              greeting = NOvlc()         
       elif 12 <= righnow.hour < 18:                                                   
              greet    = "afternoon"                                                 
              greeting = NOvlc()
       else:                                                                      
              greet    = "evening"                                                    
              greeting = NOvlc()
os.chdir(here)#---------------------------------for some weird vlc module reason, our location changes to the vlc folder after trying to import the vlc module (which currently fails). This code fixes that


# defining AJ function
def AJ():
       effect1.stop()
       effect1.play()
       newentry = ""
       print(clear)
       print("\n"*10)
       print("""






                                                                                                                  
                                                                                                         """+Fore.LIGHTRED_EX+"""   ##############      """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""    ##################  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""  #              #     """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+""" #                #    """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""####################   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""####################   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""   #                 #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""    #                #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""     #              #   """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""      #            #    """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #  """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""#       ############     #"""+Style.RESET_ALL+Style.BRIGHT+"""                                                         """+Fore.LIGHTBLACK_EX+"""["""+Fore.LIGHTRED_EX+"""A"""+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""utomated """+Fore.LIGHTRED_EX+"""J"""+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""ournal] ("""+Fore.LIGHTRED_EX+"""v7.0"""+Style.RESET_ALL+Fore.LIGHTBLACK_EX+""")"""+Style.RESET_ALL+"""
|-------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------|""")
       print("")
       print(now+Style.RESET_ALL)
       newentry = input("")
       if   os.path.exists("journal.docx") == True:
              file              = docx.Document("journal.docx")
              AJ1               = file.add_paragraph().add_run(now)
              AJ2               = file.add_paragraph()
              AJ3               = AJ2.add_run(newentry) 
              AJ1.bold          = True
              AJ3.add_break()
              file.save("journal.docx")
       elif os.path.exists("journal.docx") == False:
              file              = docx.Document()
              AJ0               = file.add_paragraph().add_run("                     AJ")
              AJ0.font.size     = Pt(41)
              AJ0.bold          = True
              AJ0run            = file.add_paragraph().add_run("______________________________________________")
              AJ0run.font.size  = Pt(25)
              AJ0run.bold       = True
              AJ1               = file.add_paragraph().add_run()
              AJ1.add_break()
              AJ1.add_break()
              AJ2               = file.add_paragraph().add_run(now)
              AJ3               = file.add_paragraph().add_run(newentry) 
              AJ2.bold          = True
              AJ3.add_break()
              file.save("journal.docx")
       print(clear)
       print("\n"*20)
       print(Fore.LIGHTGREEN_EX+"")
       tprint("       AJ ENTRY SAVED!", font="cyber")
       good3.stop()
       good3.play()
       if os.path.exists('log.txt') == False:
           log = open('log.txt','w')
           log.write(now + " |  New AJ-Entry.")
           log.close()
       elif os.path.exists('log.txt') == True:
           log = open('log.txt','a')
           log.write("\n" + now + " |  New AJ-Entry.")
           log.close()
       print(""+Style.RESET_ALL)
       time.sleep(0.5)
       os._exit(0)

AJ()
